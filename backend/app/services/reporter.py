from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
import io
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ReportGenerator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        # PDF Styles
        self.styles.add(ParagraphStyle(
            name='TitlePrimary',
            parent=self.styles['Title'],
            fontSize=28,
            spaceAfter=40,
            textColor=colors.HexColor("#1e3a8a"),  # Blue 900
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ))
        self.styles.add(ParagraphStyle(
            name='SectionHeader',
            fontSize=18,
            spaceBefore=30,
            spaceAfter=20,
            textColor=colors.HexColor("#2563eb"),  # Blue 600
            fontName='Helvetica-Bold',
            borderPadding=(0, 0, 5, 0),
            borderWidth=0,
            borderColor=colors.HexColor("#2563eb")
        ))
        self.styles.add(ParagraphStyle(
            name='HallazgoTitle',
            fontSize=13,
            spaceBefore=15,
            spaceAfter=10,
            fontName='Helvetica-Bold',
            textColor=colors.HexColor("#374151")  # Gray 700
        ))
        self.styles.add(ParagraphStyle(
            name='EvidenceStyle',
            fontSize=10,
            leftIndent=25,
            rightIndent=25,
            spaceBefore=5,
            spaceAfter=5,
            fontName='Helvetica-Oblique',
            textColor=colors.HexColor("#6b7280"),  # Gray 500
            backColor=colors.HexColor("#f9fafb"),
            borderPadding=10,
            borderRadius=5
        ))
        self.styles.add(ParagraphStyle(
            name='NormalJustified',
            parent=self.styles['Normal'],
            alignment=TA_JUSTIFY,
            fontSize=11,
            leading=14,
            spaceAfter=10
        ))

    def generate_pdf(self, data: dict) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter, 
            rightMargin=50, 
            leftMargin=50, 
            topMargin=50, 
            bottomMargin=50
        )
        elements = []

        # --- Portada ---
        elements.append(Spacer(1, 100))
        elements.append(Paragraph("INFORME DE AUDITORÍA PEDAGÓGICA", self.styles['TitlePrimary']))
        elements.append(Spacer(1, 20))
        elements.append(Paragraph(f"<b>Institución:</b> {data['institution_info']['name']}", self.styles['Heading2']))
        elements.append(Paragraph(f"<b>Tutor Responsable:</b> {data['institution_info']['tutor']}", self.styles['Normal']))
        elements.append(Spacer(1, 200))
        elements.append(Paragraph("Generado por PTAFI-AI v2.0", self.styles['Normal']))
        elements.append(PageBreak())

        # --- 1. Matriz de Hallazgos ---
        elements.append(Paragraph("1. MATRIZ DE SISTEMATIZACIÓN DE HALLAZGOS", self.styles['SectionHeader']))
        
        for item in data.get('matrix', []):
            elements.append(Paragraph(f"Categoría: {item['category_name'].upper()}", self.styles['HallazgoTitle']))
            elements.append(Paragraph(f"<b>Hallazgo:</b> {item['hallazgo']}", self.styles['NormalJustified']))
            
            evid = item.get('evidencia', {})
            elements.append(Paragraph(
                f"<i>Evidencia:</i> \"{evid.get('text', 'N/A')}\"<br/>"
                f"<font size='8' color='gray'>(Documento: {evid.get('document_name', 'N/A')} · Pág: {evid.get('page', 'N/A')})</font>", 
                self.styles['EvidenceStyle']
            ))
            
            elements.append(Paragraph(f"<b>Interpretación Pedagógica:</b> {item['interpretacion']}", self.styles['NormalJustified']))
            elements.append(Paragraph(f"<b>Implicación para el PFI:</b> {item['implicacion_pfi']}", self.styles['NormalJustified']))
            elements.append(Spacer(1, 15))

        elements.append(PageBreak())

        # --- 2. Reporte de Calidad (Pilares) ---
        elements.append(Paragraph("2. REPORTE DE CALIDAD (PILARES ESTRATÉGICOS)", self.styles['SectionHeader']))
        
        for p in data.get('quality_report', []):
            elements.append(Paragraph(f"{p['pillar_name']} - Evaluación: {p['score']}/10", self.styles['HallazgoTitle']))
            elements.append(Paragraph(f"<b>Análisis Crítico:</b> {p['analysis']}", self.styles['NormalJustified']))
            
            recs = p.get('recommendations', [])
            if recs:
                elements.append(Paragraph("<b>Recomendaciones de Mejora:</b>", self.styles['Normal']))
                for r in recs:
                    elements.append(Paragraph(f"• {r}", self.styles['Normal'], bulletText='•'))
            elements.append(Spacer(1, 15))

        doc.build(elements)
        pdf_value = buffer.getvalue()
        buffer.close()
        return pdf_value

    def generate_docx(self, data: dict) -> bytes:
        doc = Document()
        
        # --- Estilos DOCX ---
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # --- Portada ---
        title = doc.add_heading('INFORME DE AUDITORÍA PEDAGÓGICA', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_spacer = lambda: doc.add_paragraph().add_run().add_break()
        for _ in range(3): doc.add_paragraph()
        
        inst_p = doc.add_paragraph()
        inst_p.add_run('Institución: ').bold = True
        inst_p.add_run(data['institution_info']['name'])
        inst_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tutor_p = doc.add_paragraph()
        tutor_p.add_run('Tutor Responsable: ').bold = True
        tutor_p.add_run(data['institution_info']['tutor'])
        tutor_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # --- 1. Matriz ---
        doc.add_heading('1. MATRIZ DE SISTEMATIZACIÓN DE HALLAZGOS', level=1)
        
        for item in data.get('matrix', []):
            doc.add_heading(f"Categoría: {item['category_name']}", level=2)
            
            p = doc.add_paragraph()
            p.add_run('Hallazgo: ').bold = True
            p.add_run(item['hallazgo'])
            
            # Evidencia en un cuadro (tabla simple de 1 sola celda)
            evid = item.get('evidencia', {})
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Light Shading Accent 1'
            cell = table.rows[0].cells[0]
            cell.text = f"\"{evid.get('text', 'N/A')}\" \n(Fuente: {evid.get('document_name', 'N/A')} - Pág {evid.get('page', 'N/A')})"
            
            doc.add_paragraph() # Spacer
            
            interp = doc.add_paragraph()
            interp.add_run('Interpretación Pedagógica: ').bold = True
            interp.add_run(item['interpretacion'])
            
            impl = doc.add_paragraph()
            impl.add_run('Implicación PFI: ').bold = True
            impl.add_run(item['implicacion_pfi'])
            
            doc.add_paragraph('---')

        doc.add_page_break()

        # --- 2. Calidad ---
        doc.add_heading('2. REPORTE DE CALIDAD (PILARES ESTRATÉGICOS)', level=1)
        
        for p in data.get('quality_report', []):
            doc.add_heading(f"{p['pillar_name']} (Puntaje: {p['score']}/10)", level=2)
            
            analysis = doc.add_paragraph()
            analysis.add_run('Análisis: ').bold = True
            analysis.add_run(p['analysis'])
            
            recs = p.get('recommendations', [])
            if recs:
                doc.add_paragraph('Recomendaciones:', style='List Bullet')
                for r in recs:
                    doc.add_paragraph(r, style='List Bullet 2')
            
            doc.add_paragraph()

        # Guardar a buffer
        target = io.BytesIO()
        doc.save(target)
        docx_value = target.getvalue()
        target.close()
        return docx_value

pdf_reporter = ReportGenerator()

